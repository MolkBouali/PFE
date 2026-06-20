"""
Service generation avis officiel PDF.
- Recupere les donnees du dossier depuis la base de donnees
- Determine le type d avis selon les validations DER/DTA/DANA/DNA
- Remplit le template Word officiel OACA avec les donnees
- Sauvegarde le fichier et enregistre le document en base
"""
import os, datetime
from sqlalchemy.orm import Session
from docx import Document
from backend.models.dossier import Dossier
from backend.models.document_genere import DocumentGenere
from backend.schemas.decision_schema import AvisValidationRequest
from backend.core.config import settings

class PDFService:
    def __init__(self, db: Session):
        self.db = db

    def generate(self, dossier_id: int, data: AvisValidationRequest, agent_id: int) -> str:
        dossier = self.db.query(Dossier).filter(Dossier.id == dossier_id).first()
        dirs_valides = sum([data.validation_der, data.validation_dta,
                            data.validation_dana, data.validation_dna])
        type_avis = "FAVORABLE" if dirs_valides >= 3 else "DEFAVORABLE"

        template_path = os.path.join(settings.TEMPLATES_PATH, "avis_template.docx")
        if not os.path.exists(template_path):
            from docx import Document as D
            doc = D()
        else:
            doc = Document(template_path)

        placeholders = {"{{NUMERO_DOSSIER}}": dossier.numero_dossier,
                        "{{NOM_DEMANDEUR}}": dossier.nom_demandeur,
                        "{{DATE}}": datetime.date.today().strftime("%d/%m/%Y"),
                        "{{TYPE_AVIS}}": type_avis,
                        "{{JUSTIFICATION}}": data.justification or ""}
        for para in doc.paragraphs:
            for k, v in placeholders.items():
                if k in para.text:
                    para.text = para.text.replace(k, v)

        os.makedirs(os.path.join(settings.STORAGE_PATH, "pdf_avis"), exist_ok=True)
        output_path = os.path.join(settings.STORAGE_PATH, "pdf_avis",
                                   f"avis_{dossier_id}_{datetime.date.today()}.docx")
        doc.save(output_path)

        doc_db = DocumentGenere(dossier_id=dossier_id, nom_fichier=os.path.basename(output_path),
                                type_document="AVIS_PDF", chemin_stockage=output_path,
                                date_creation=datetime.datetime.now())
        self.db.add(doc_db)
        self.db.commit()
        return output_path

    def generate_complement(self, dossier_id: int) -> str:
        """Génère le document de demande de complément de coordonnées au format DMS."""
        dossier = self.db.query(Dossier).filter(Dossier.id == dossier_id).first()
        if not dossier:
            raise ValueError("Dossier non trouvé")

        template_path = os.path.join(settings.TEMPLATES_PATH, "complement_template.docx")
        
        if not os.path.exists(template_path):
            # Création du dossier et d'un template minimal si absent
            os.makedirs(settings.TEMPLATES_PATH, exist_ok=True)
            doc = Document()
            doc.add_heading('DEMANDE DE COMPLÉMENT DE COORDONNÉES', 0)
            doc.add_paragraph('\n{{NUMERO_DOSSIER}}\n{{NOM_DEMANDEUR}}')
            doc.add_paragraph('\nCher demandeur,')
            doc.add_paragraph('\nAprès analyse de votre dossier, nous avons constaté que les coordonnées fournies ne sont pas au format requis.')
            doc.add_paragraph('\nNous vous prions de bien vouloir compléter votre dossier en fournissant des coordonnées qui respectent strictement le format DMS (Degrés, Minutes, Secondes).')
            doc.add_paragraph('\nCordialement,\nL\'OACA')
            doc.save(template_path)
        else:
            doc = Document(template_path)

        placeholders = {
            "{{NUMERO_DOSSIER}}": f"Dossier N°: {dossier.numero_dossier}",
            "{{NOM_DEMANDEUR}}": f"Nom: {dossier.nom_demandeur}",
            "{{DATE}}": datetime.date.today().strftime("%d/%m/%Y"),
        }
        
        for para in doc.paragraphs:
            for k, v in placeholders.items():
                if k in para.text:
                    para.text = para.text.replace(k, v)

        os.makedirs(os.path.join(settings.STORAGE_PATH, "pdf_complement"), exist_ok=True)
        output_path = os.path.join(settings.STORAGE_PATH, "pdf_complement",
                                    f"complement_{dossier_id}_{datetime.date.today()}.docx")
        doc.save(output_path)

        doc_db = DocumentGenere(
            dossier_id=dossier_id, 
            nom_fichier=os.path.basename(output_path),
            type_document="COMPLEMENT_DMS", 
            chemin_stockage=output_path,
            date_creation=datetime.datetime.now()
        )
        self.db.add(doc_db)
        self.db.commit()
        return output_path